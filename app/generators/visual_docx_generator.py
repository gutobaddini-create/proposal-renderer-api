from __future__ import annotations

from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from docx.shared import Cm


def generate_visual_docx_from_pdf(pdf_path: Path, output_dir: Path) -> Path:
    path = output_dir / "proposta_visual.docx"
    image_dir = output_dir / "_visual_pages"
    image_dir.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    pdf = pdfium.PdfDocument(str(pdf_path))
    try:
        for index in range(len(pdf)):
            page = pdf[index]
            bitmap = page.render(scale=2.0)
            image = bitmap.to_pil()
            image_path = image_dir / f"page_{index + 1}.png"
            image.save(image_path)

            if index:
                document.add_page_break()
            document.add_picture(str(image_path), width=Cm(21))
    finally:
        pdf.close()

    document.save(path)
    return path
