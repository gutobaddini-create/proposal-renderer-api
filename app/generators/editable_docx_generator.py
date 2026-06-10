from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.brand_assets import BRAND_BLUE, get_logo_path


BRAND_BLUE_HEX = BRAND_BLUE.replace("#", "")
NAVY = RGBColor(16, 17, 48)
GOLD = RGBColor(181, 138, 69)
MUTED = RGBColor(102, 112, 133)


def _clean(value: Any) -> str:
    return str(value or "").strip()


def _add_key_value(document: Document, label: str, value: Any) -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = NAVY
    paragraph.add_run(_clean(value))


def _set_document_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = NAVY
        style.font.bold = True


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_text_color(cell, color: RGBColor) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color


def _apply_header_footer(document: Document) -> None:
    logo_path = get_logo_path()
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.45)

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Cm(17))
    header_table.autofit = False
    header_table.columns[0].width = Cm(8)
    header_table.columns[1].width = Cm(9)

    logo_paragraph = header_table.cell(0, 0).paragraphs[0]
    logo_paragraph.add_run().add_picture(str(logo_path), width=Cm(4.2))

    text_paragraph = header_table.cell(0, 1).paragraphs[0]
    text_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = text_paragraph.add_run("Proposta juridica empresarial")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(9)

    for cell in header_table.rows[0].cells:
        _set_cell_shading(cell, BRAND_BLUE_HEX)

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Cm(17))
    footer_table.autofit = False
    footer_table.columns[0].width = Cm(12)
    footer_table.columns[1].width = Cm(5)
    left = footer_table.cell(0, 0).paragraphs[0]
    left.add_run("Lobo Baddini Advocacia").bold = True
    left.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    right = footer_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right.add_run("Documento confidencial")
    right_run.font.color.rgb = RGBColor(255, 255, 255)
    right_run.font.size = Pt(8)

    for cell in footer_table.rows[0].cells:
        _set_cell_shading(cell, BRAND_BLUE_HEX)
        _set_cell_text_color(cell, RGBColor(255, 255, 255))


def _add_section_content(document: Document, content: Any) -> None:
    if isinstance(content, list):
        for item in content:
            document.add_paragraph(_clean(item), style="List Bullet")
    elif isinstance(content, dict):
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for key, value in content.items():
            cells = table.add_row().cells
            cells[0].text = _clean(key)
            cells[1].text = _clean(value)
    else:
        document.add_paragraph(_clean(content))


def generate_editable_docx(data: dict[str, Any], output_dir: Path) -> Path:
    path = output_dir / "proposta_editavel.docx"
    document = Document()
    _set_document_styles(document)
    _apply_header_footer(document)

    provider = data.get("provider", {})
    client = data.get("client", {})
    proposal = data.get("proposal", {})
    investment = proposal.get("investment", {})
    term = proposal.get("term", {})

    title = document.add_heading(proposal.get("title", "Proposta Jurídica"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        title.runs[0].font.color.rgb = NAVY

    if proposal.get("subtitle"):
        subtitle = document.add_paragraph(proposal["subtitle"])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.color.rgb = MUTED

    document.add_paragraph()
    document.add_heading("Dados da proposta", level=1)
    _add_key_value(document, "Contratante", client.get("name"))
    _add_key_value(document, "Proponente", provider.get("name"))
    _add_key_value(document, "Responsável", provider.get("responsible"))
    _add_key_value(document, "E-mail", provider.get("email"))

    if proposal.get("objective"):
        document.add_heading("Objetivo", level=1)
        document.add_paragraph(_clean(proposal.get("objective")))

    document.add_heading("Investimento", level=1)
    document.add_paragraph(_clean(investment.get("amountText", "")))

    document.add_heading("Prazo", level=1)
    document.add_paragraph(_clean(term.get("duration", "")))

    document.add_heading("Escopo", level=1)
    for section in data.get("layoutPlan", []):
        document.add_heading(_clean(section.get("title") or section.get("heading") or "Seção"), level=2)
        content = section.get("content") or section.get("body") or section.get("text") or section
        _add_section_content(document, content)

    document.add_heading("Aceite", level=1)
    document.add_paragraph("De acordo com os termos apresentados nesta proposta.")
    document.add_paragraph("\nContratante: ______________________________")
    document.add_paragraph("Proponente: ______________________________")

    document.save(path)
    return path
